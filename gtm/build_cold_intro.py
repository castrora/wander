"""
Wander — Cold Introduction Email Builder
Generates gtm/wander-cold-intro-emails.docx

Four audience-specific versions:
  A. Port of Seattle / Terminal Access
  B. Tourism Organization (Visit Seattle / State Tourism)
  C. Cruise Line (Shore Excursions / Partnerships)
  D. Hotel (GM / Director of Sales)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)

# ── Colour constants ─────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
TEAL   = RGBColor(0x1A, 0x7A, 0x6E)
GREY   = RGBColor(0x59, 0x59, 0x59)
BLACK  = RGBColor(0x00, 0x00, 0x00)
ORANGE = RGBColor(0xC5, 0x5A, 0x11)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_para(doc, text="", style="Normal", bold=False, italic=False,
             size=11, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return p

def add_rule(doc, color_hex="1A7A6E"):
    """Thin horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_label(doc, letter, title, subtitle):
    """Section header block for each email version."""
    add_para(doc, space_before=14, space_after=0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(f"VERSION {letter}  |  ")
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = NAVY; r1.font.name = "Calibri"
    r2 = p.add_run(title)
    r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = TEAL; r2.font.name = "Calibri"
    add_para(doc, subtitle, italic=True, size=10, color=GREY, space_before=1, space_after=4)
    add_rule(doc, "1A7A6E")
    add_para(doc, space_before=2, space_after=0)

def add_field(doc, label, value):
    """Email field row: To / Subject / etc."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY; r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(10); r2.font.color.rgb = BLACK; r2.font.name = "Calibri"

def add_body(doc, paragraphs):
    """Body text — list of strings, each becomes a paragraph."""
    add_para(doc, space_before=6, space_after=0)
    for i, text in enumerate(paragraphs):
        space_after = 7 if i < len(paragraphs) - 1 else 4
        add_para(doc, text, size=10.5, color=BLACK,
                 space_before=0, space_after=space_after)

def add_signature(doc):
    add_para(doc, space_before=4, space_after=2)
    add_para(doc, "Ray Castro", bold=True, size=10.5, color=NAVY, space_before=0, space_after=1)
    add_para(doc, "Founder, Wander", size=10, color=GREY, space_before=0, space_after=1)
    add_para(doc, "ray.s.castro@outlook.com", size=10, color=TEAL, space_before=0, space_after=1)
    add_para(doc, "pierway.io  ·  Seattle, WA", italic=True, size=9.5, color=GREY,
             space_before=0, space_after=0)

def add_note(doc, text):
    """Editorial note in brackets — for Ray's reference, not sent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"✎  {text}")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = ORANGE; r.font.name = "Calibri"

def page_break(doc):
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ════════════════════════════════════════════════════════════════════════════
add_para(doc, "WANDER", bold=True, size=28, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=2)
add_para(doc, "Cold Introduction Emails — 2026 Cruise Season Outreach",
         size=12, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=4)
add_rule(doc)
add_para(doc,
    "Four audience-specific versions. Personalize the [BRACKETS] before sending. "
    "Keep each email under 200 words — these are busy people.",
    italic=True, size=10, color=GREY,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)

versions_table = [
    ("A", "Port of Seattle",         "Marie Ellingson — Manager, Cruise Services & Business Development"),
    ("B", "Tourism Organization",    "Marco Leal — VP Destination Development, Visit Seattle"),
    ("C", "Cruise Line",             "Koreen McNutt — VP Business Development, Virgin Voyages"),
    ("D", "Hotel",                   "Haley Connors — Destination Sales Executive, Marriott Waterfront"),
]
add_para(doc, space_before=6, space_after=2)
for letter, audience, example in versions_table:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  Version {letter}  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY; r1.font.name = "Calibri"
    r2 = p.add_run(f"→  {audience}  ")
    r2.font.size = Pt(10); r2.font.color.rgb = BLACK; r2.font.name = "Calibri"
    r3 = p.add_run(f"(e.g. {example})")
    r3.italic = True; r3.font.size = Pt(9.5); r3.font.color.rgb = GREY; r3.font.name = "Calibri"

page_break(doc)


# ════════════════════════════════════════════════════════════════════════════
# VERSION A — PORT OF SEATTLE
# ════════════════════════════════════════════════════════════════════════════
add_label(doc, "A", "Port of Seattle",
          "Best for: Marie Ellingson (mellingson@portseattle.org) · Rosie Courtney")

add_field(doc, "To",      "[Marie Ellingson]  <mellingson@portseattle.org>")
add_field(doc, "Subject", "Wander — AI Walking Experience for Seattle Cruise Passengers | QR Code Placement Inquiry")
add_field(doc, "From",    "Ray Castro  <ray.s.castro@outlook.com>")

add_body(doc, [
    "Hi Marie,",

    "I'm Ray Castro, founder of Wander — an AI-powered walking experience platform launching in Seattle for the 2026 cruise season. Wander generates personalized walking routes for cruise passengers during their Seattle port days, embedding contextual recommendations for local merchants along the way.",

    "I'm reaching out because I'd like to understand the process for placing Wander QR codes at Pier 91 and Pier 66 — either at the disembarkation area, the visitor information point, or through a partner at the terminal. The goal is to give passengers a free, immediately useful resource the moment they step off the ship.",

    "The model is simple: passengers scan, get a personalized walking route, and discover Seattle on their own terms. Merchants pay Wander only when a passenger actually walks through their door — no impressions, no clicks, just verified visits.",

    "Would you be open to a 20-minute call this week or next to understand what's possible? I'm flexible on timing and happy to come to the port in person.",
])

add_signature(doc)
add_note(doc, "Personalize: add a specific reference to Port of Seattle if you find a recent news item or initiative — e.g. their sustainability program or 2026 season announcements.")

page_break(doc)


# ════════════════════════════════════════════════════════════════════════════
# VERSION B — TOURISM ORGANIZATION
# ════════════════════════════════════════════════════════════════════════════
add_label(doc, "B", "Tourism Organization",
          "Best for: Marco Leal (mleal@visitseattle.org) · Tammy Canavan · State of WA Tourism")

add_field(doc, "To",      "[Marco Leal]  <mleal@visitseattle.org>")
add_field(doc, "Subject", "Wander — AI Walking Platform for Seattle's 2M+ Cruise Visitors in 2026 | Partnership Inquiry")
add_field(doc, "From",    "Ray Castro  <ray.s.castro@outlook.com>")

add_body(doc, [
    "Hi Marco,",

    "I'm Ray Castro, founder of Wander — an AI-powered urban walking platform built to convert Seattle's 2M+ annual cruise passengers into active local explorers and verified commerce for Seattle merchants.",

    "Wander generates real-time, personalized walking routes for passengers during their Seattle port days — embedding contextual storytelling about the city alongside optional local merchant recommendations. Merchants pay only for confirmed in-store visits. No impressions. No ads. Just performance-based local commerce.",

    "Given Visit Seattle's mandate around destination development and the record 2026 cruise season ahead, I'd love to explore what a partnership could look like — whether that's QR code placement through your hotel network, co-distribution through tourism partners, or something else.",

    "I'd welcome 20 minutes at your convenience. Happy to share a one-pager on the model beforehand if that's useful.",
])

add_signature(doc)
add_note(doc, "For Tammy Canavan: adjust the subject line to reference the FIFA World Cup 2026 opportunity as well — Visit Seattle is actively working on that and it shows you've done your homework.")

page_break(doc)


# ════════════════════════════════════════════════════════════════════════════
# VERSION C — CRUISE LINE
# ════════════════════════════════════════════════════════════════════════════
add_label(doc, "C", "Cruise Line — Shore Excursions / Partnerships",
          "Best for: Koreen McNutt (Virgin Voyages) · Katty Byrd (NCL) · Jessica Ashe (Holland America)")

add_field(doc, "To",      "[Koreen McNutt]  <[First.Last@virginvoyages.com]>")
add_field(doc, "Subject", "Wander + [Virgin Voyages] — AI Walking Experience for Seattle Port Days | 2026 Season")
add_field(doc, "From",    "Ray Castro  <ray.s.castro@outlook.com>")

add_body(doc, [
    "Hi Koreen,",

    "Congratulations on bringing Brilliant Lady to Seattle — it's an exciting debut for the Alaska season.",

    "I'm Ray Castro, founder of Wander. We're building an AI-powered walking experience for cruise passengers during Seattle port days — personalized routes that help Sailors discover the city on their own terms, with optional merchant recommendations woven into the narrative. No sponsored stops. No ads. Just genuinely useful exploration.",

    "The model is performance-based: merchants pay Wander only for confirmed visits, so there's no cost to [Virgin Voyages] and no obligation on Sailors. It's simply a better port day option — available through a QR code at disembarkation or in your pre-departure communications.",

    "Given that you're standing up your Seattle program right now, I wanted to reach out early. Would you be open to a brief call this month to explore whether Wander fits within [Virgin Voyages]'s Seattle experience?",
])

add_signature(doc)
add_note(doc, "Customize the [bracketed] fields for each cruise line. For NCL's Katty Byrd, drop the 'congratulations' opener. For HAL's Jessica Ashe, reference their '80 years in Alaska' milestone in the subject line.")

page_break(doc)


# ════════════════════════════════════════════════════════════════════════════
# VERSION D — HOTEL
# ════════════════════════════════════════════════════════════════════════════
add_label(doc, "D", "Hotel — GM / Director of Sales",
          "Best for: Haley Connors (Marriott Waterfront) · Ian McClendon (Edgewater) · Leslie Womack (Mayflower)")

add_field(doc, "To",      "[Haley Connors]  <[direct or hotel general line]>")
add_field(doc, "Subject", "Wander — Complimentary Port Day Experience for [Marriott Waterfront] Cruise Guests | 2026 Season")
add_field(doc, "From",    "Ray Castro  <ray.s.castro@outlook.com>")

add_body(doc, [
    "Hi Haley,",

    "I'm Ray Castro, founder of Wander — an AI-powered walking experience platform launching in Seattle for the 2026 cruise season.",

    "Wander gives cruise passengers a personalized walking route for their Seattle port day — AI-generated, real-time, and built around their available time and interests. It's a genuinely useful tool for guests who want to explore the city independently rather than book a bus tour.",

    "I'd love to explore a simple partnership: a Wander QR code at the [Marriott Waterfront] concierge desk, front desk, or in-room materials for cruise package guests. No cost to the hotel, no integration required — just a better resource for your guests the morning they head to the port.",

    "Given [Marriott Waterfront]'s Park & Cruise package and your location steps from Pier 66, your guests are exactly who Wander is built for. Would you be open to a quick call or coffee to see if it's a fit?",
])

add_signature(doc)
add_note(doc, "For Edgewater: mention their recent property refresh and waterfront renovation — shows you know the property. For Mayflower: reference their shuttle program to both Pier 66 and Pier 91 as the distribution context.")

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"C:\Users\raysc\wander\gtm\wander-cold-intro-emails.docx"
doc.save(out)
print(f"Saved: {out}")
